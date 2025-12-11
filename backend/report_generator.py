import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def generate_passport_report(passport_data):
    document = Document()
    
    # Basic Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Header ---
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПЕРЕВОД С АНГЛИЙСКОГО ЯЗЫКА НА РУССКИЙ ЯЗЫК")
    run.bold = True
    run.font.size = Pt(14)
    
    document.add_paragraph() # Spacer

    # --- Bio Page ---
    document.add_paragraph("[Страница с персональными данными]")
    
    bio = passport_data.get('biographical_page', {})
    nationality = bio.get('nationality', '')
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"ПАСПОРТ {nationality}")
    run.bold = True
    run.font.size = Pt(14)

    # Bio Fields List
    def add_line(label, value):
        if not value: return
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(label + ": ").bold = True
        p.add_run(str(value).upper())

    add_line("Тип", "P")
    nationality_code = bio.get('nationality') or ""
    add_line("Код государства", nationality_code[:3] if nationality_code else "")
    add_line("Номер паспорта", bio.get('passport_number'))
    add_line("Фамилия", bio.get('surname') or bio.get('full_name'))
    add_line("Имя", bio.get('given_names'))
    add_line("Гражданство", bio.get('nationality'))
    add_line("Дата рождения", bio.get('date_of_birth'))
    add_line("Пол", bio.get('gender'))
    add_line("Место рождения", bio.get('place_of_birth'))
    add_line("Дата выдачи", bio.get('issue_date'))
    add_line("Действителен до", bio.get('expiry_date'))
    add_line("Орган выдачи", bio.get('issuing_authority'))

    # MRZ
    mrz = passport_data.get('mrz', {})
    if mrz:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.add_run("Машиночитаемая зона:").bold = True
        p.add_run("\n" + (mrz.get('mrz_line1') or "") + "\n" + (mrz.get('mrz_line2') or ""))

    # --- Visas ---
    visas = passport_data.get('visas') or []
    # Ensure visas is a list and filter out None entries
    if not isinstance(visas, list):
        visas = []
    visas = [v for v in visas if v and isinstance(v, dict)]
    # Sort by page number
    visas.sort(key=lambda x: x.get('page_number') if isinstance(x.get('page_number'), int) else 999)
    
    if visas:
        for visa in visas:
            page_num = visa.get('page_number')
            page_header = f"[Стр. {page_num}: Виза]" if page_num else "[Виза]"
            document.add_paragraph("\n" + page_header)
            
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 1.2
            
            p.add_run(f"ВИЗА {visa.get('visa_number', '')}\n").bold = True
            p.add_run(f"ДЕЙСТВИТЕЛЬНА ДЛЯ: {visa.get('country', '')}\n")
            p.add_run(f"С: {visa.get('issue_date', '')}   ДО: {visa.get('expiry_date', '')}\n")
            p.add_run(f"СРОК ПРЕБЫВАНИЯ: {visa.get('stay_duration', '')}\n")
            p.add_run(f"ТИП ВИЗЫ: {visa.get('visa_type', '')}   КОЛИЧЕСТВО ВЪЕЗДОВ: {visa.get('entries_allowed', '')}\n")
            p.add_run(f"ВЫДАНО В: {visa.get('place_of_issue', '')}   ДАТА: {visa.get('issue_date', '')}\n")
            if visa.get('remarks'):
                p.add_run(f"ОТМЕТКИ: {visa.get('remarks', '')}\n")
            
            if visa.get('mrz_line1') or visa.get('mrz_line2'):
                p.add_run("\n" + (visa.get('mrz_line1') or "") + "\n" + (visa.get('mrz_line2') or ""))

    # --- Stamps ---
    stamps = passport_data.get('stamps') or []
    if not isinstance(stamps, list):
        stamps = []
    stamps = [s for s in stamps if s and isinstance(s, dict)]
    if stamps:
        document.add_paragraph("\n[Отметки о пересечении границы]")
        p = document.add_paragraph()
        
        # Simply list them as text lines as seen in sample
        for stamp in stamps:
            country = stamp.get('country', '')
            date = stamp.get('date', '')
            st_type = stamp.get('type', '')
            # Format: "Штамп: [Страна] [Дата] [Тип]"
            p.add_run(f"Штамп: {country} {date} {st_type}\n")

    document.add_paragraph("\n")
    
    # --- Footer / Certification ---
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("_" * 80)
    
    p = document.add_paragraph()
    p.add_run("Перевод выполнен переводчиком с английского языка на русский язык.\n")
    p.add_run("Я подтверждаю верность выполненного мной перевода.\n\n")
    
    p = document.add_paragraph()
    p.add_run("Переводчик: _________________________ (Подпись)")

    # Save
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return f
